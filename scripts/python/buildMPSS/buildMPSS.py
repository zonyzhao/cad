#!/usr/bin/env python

import sys
import os
import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import subprocess

warnings.filterwarnings("ignore")
#redirect error messages to stdout so cadence can read them
sys.stderr = sys.stdout

#mpss document will be outputted to the release folder of the cell design directory
# ex. ~/deslibs/ln0705/release/ln0705_0_revA.docx

#THINGS TO KEEP UP TO DATE:
#process types for GaN and GaAs, so we know which RSS sheet to look at, (directly below this)
#the location and names of the RSS files and also the sheetname of the main page with all of the information (directly below this)
#the location of the photo of the north south diagram(directly below this
#abbrevations if new layers are added (in the function getAbbrev)


#processes for each of the types, add to these if there are more proccesses
GaN =["P80A", "P80B", "P81"]
GaAs = ["P46", "P51", "P60", "P70"]
other = ["P82", "D83", "D82"]
GaNFile = "GaN_RSS.xlsx"
GaAsFile = "GaAs_RSS.xlsx"
otherFile = "other.xlsx"
GaNSheet = "GaN"
GaAsSheet = "GaAs"
otherSheet = "DEV"
rssPath = "/nfs/foundry/met/cadence/cad51Setup/"
northSouthDiagram = "/nfs/foundry/met/cadence/cad51Setup/northSouthDiagram.png"



class Layer(object):
    def __init__(self):
        #initalize
        self.mpss = None
        self.itemNo = None
        self.pLayer = None
        self.rTitle = None
        self.gds = None
        self.mag = None
        self.rField = None
        self.bias = None
        self.address = None
        self.cad = None
        self.rCD = None
        self.optional = None
    def addData(self, dataType, data):
        if dataType == "MPSS PAGE": 
            self.mpss = data
        elif dataType == "ITEM NUMBER": 
            self.itemNo = data
        elif dataType == "PROCESS LAYER":
            self.pLayer = data
        elif dataType == "RETICLE TITLE":
            self.rTitle = data
        elif dataType == "GDS LAYER":
            self.gds = data
        elif dataType == "MAG":
            self.mag = data
        elif dataType == "RETICLE FIELD":
            self.rField = data
        elif dataType == "BIAS/SIDE":
            self.bias = data
        elif dataType == "ADDRESS":
            self.address = data
        elif dataType == "CAD DATA CD":
            self.cad = data
        elif dataType == "RETICLE CD/TOLERANCE":
            self.rCD = data
        else:
            self.optional = data
    def getData(self, dataType):
        if dataType == "MPSS PAGE": 
            data =  self.mpss
        elif dataType == "ITEM NUMBER": 
            data = self.itemNo
        elif dataType == "PROCESS LAYER":
            data = self.pLayer
        elif dataType == "RETICLE TITLE":
            data = self.rTitle
        elif dataType == "GDS LAYER":
            data = self.gds
        elif dataType == "MAG":
            data = self.mag
        elif dataType == "RETICLE FIELD":
            data = self.rField
        elif dataType == "BIAS/SIDE":
            data = self.bias
        elif dataType == "ADDRESS":
            data = self.address
        elif dataType == "CAD DATA CD":
            data = self.cad
        elif dataType == "RETICLE CD/TOLERANCE":
            data = self.rCD
        else:
            data = self.optional
        return str(data)


def rowEmpty(worksheet, row):
    for col in range(1, worksheet.max_column+1):
        if worksheet.cell(row=row, column=col).value != None:
            return False
    return True


def process_worksheet(process):
    ws = None
    filename = None
    sheetName = None
    #get the excel sheet for the process type
    if process in GaN:
        filename = GaNFile
        sheetName = GaNSheet
    elif process in GaAs:
        filename = GaAsFile
        sheetName = GaAsSheet
    elif process in other:
        filename = otherFile
        sheetName = otherSheet
    else:
        print("Process \"%s\" not recognized, aborting..." % process)
        sys.exit()
    ws = load_workbook(rssPath+filename).get_sheet_by_name(sheetName)
    #data we need to pull
    categories = ["MPSS PAGE", "ITEM NUMBER", "PROCESS LAYER", "RETICLE TITLE", 
                "GDS LAYER", "MAG", "RETICLE FIELD", "BIAS/SIDE", "ADDRESS", 
                "CAD DATA CD", "RETICLE CD/TOLERANCE"]
    specs = {}
    #map the process and each of the categories to their column numbers
    for col in range(1, ws.max_column+1):
        cell = ws.cell(row=2, column=col) #row 2 is where all the headers are
        #strip out any leading and trailing whitepsace and ensure that all words are separated by one space for robustness
        header = None
        if cell.value != None:
            header = " ".join(cell.value.strip().split())
        if (header in categories) or cell.value == process:
            specs[cell.value] = col
    #now go through and grab data for each layer
    layerList = []
    prevValue = "N"
    for r in range(3, ws.max_row+1):
        if rowEmpty(ws, r):
            break #we finished all the layers
        else: #go through and grab the data
            newLayer = Layer()
            col = specs[process]
            cell = (ws.cell(row = r, column = col))
            value = cell.value
            if value == "N":
                prevValue = "N"
                continue
            if value == None and prevValue == "N": continue
            if value == None: value = "NA"
            newLayer.addData(process, value)
            for attribute in categories:
                col = specs[attribute]
                cell = ws.cell(row = r, column = col)
                value = cell.value
                if value == None: value = "NA"
                newLayer.addData(attribute, value)
            layerList.append(newLayer)
            prevValue = "Y"
    return layerList

def getAbbrev(title):
    #strip out any leading and trailing whitepsace and ensure that all words are separated by one space for robustness
    pureTitle = title.strip()
    pureTitle = "".join(pureTitle.split())
    abbrev = title
    if pureTitle == "ZERO": abbrev = "ZERO"
    elif pureTitle == "MESA": abbrev = "MESA"
    elif pureTitle == "SOURCEDRAINPRI": abbrev = "SD PRI"
    elif pureTitle == "SOURCEDRAINSEC": abbrev = "SD SEC"
    elif pureTitle == "TGTOP": abbrev = "TGTOP"
    elif pureTitle == "CAPBOTTOM": abbrev = "CB"
    elif pureTitle == "TANTALUMPROTECT": abbrev = "TA PRO"
    elif pureTitle == "TANTALUM": abbrev = "TA"
    elif pureTitle == "NITRIDEETCH": abbrev = "NE"
    elif pureTitle == "THICKMETALPOST": abbrev = "TH POST"
    elif pureTitle == "THICKMETALPLATE": abbrev = "TH PLATE"
    elif pureTitle == "THICKMETALPLATETRI": abbrev = "TH PLATE TRI"
    elif pureTitle == "AIRBRIDGE": abbrev = "AB"
    elif pureTitle == "THICKMETAL": abbrev = "TH"
    elif pureTitle == "GLASS": abbrev = "GLASS"
    elif pureTitle == "VIA": abbrev = "VIA"
    elif pureTitle == "SDNUMBERS": abbrev = "NUM-N"
    elif pureTitle == "SOLDERSTOP": abbrev = "SSTOP"
    elif pureTitle == "BCBPLUG": abbrev = "BCB PLUG"
    elif pureTitle == "BCB": abbrev = "BCB"
    elif pureTitle == "BCBINV": abbrev = "BCB INV"
    elif pureTitle == "BCBLID": abbrev = "BCB LID"
    elif pureTitle == "BCBCORRAL": abbrev = "BCB CORRAL"
    elif pureTitle == "NITRIDEETCH1": abbrev = "NE1"
    elif pureTitle == "THINTANTALUM": abbrev = "THIN TA"
    elif pureTitle == "SCFP": abbrev = "SCFP"
    elif pureTitle == "METAL2": abbrev = "MET2"
    elif pureTitle == "RECESS": abbrev = "REC"
    elif pureTitle == "OPTOP": abbrev = "OPTOP"
    elif pureTitle == "GATESTEM": abbrev = "GSTEM"
    elif pureTitle == "GRID": abbrev = "GRID"
    elif pureTitle == "ZEROPRI": abbrev = "ZERO PRI"
    elif pureTitle == "ZEROSEC": abbrev = "ZERO SEC"
    elif pureTitle == "SOURCEDRAIN": abbrev = "SD"
    elif pureTitle == "SOURCEDRAINTEXT": abbrev = "SD TEXT"
    return abbrev

def build_filename(layer, designName):
    abbrev = getAbbrev(layer.rTitle)
    end = abbrev.split()[-1]
    midLetter = "P"
    if end == "SEC": midLetter = "S"
    elif end == "TRI": midLetter = "T"
    elif layer.mpss == "Numbers": midLetter = "N"
    elif layer.mpss == "Backside": midLetter = "W"
    front = designName[:2].upper() + designName[3:]
    filename = "%s0%s%02d.00" % (front, midLetter, long(layer.gds))
    return filename

def get_user():
    #grabs the stdout of the cmd and strips off the trailing newline '\n' character
    name = subprocess.Popen(["whoami"], stdout = subprocess.PIPE).communicate()[0][:-1]
    user = "Unrecognized user"
    if name == "uptonl": user = "Laura Upton"
    elif name == "nardonem": user = "Marie Nardone"
    elif name == "zeidlerm": user = "Maureen Zeidler"
    elif name == "bosleyj": user == "Joe Bosley"
    elif name == "charbonneaut": user = "Tom Charbonneau"
    elif name == "tarnuzzere": user = "Ed Tarnuzzer"
    elif name == "contem": user = "Matt Conte"
    elif name == "parkesm": user = "Mike Parkes"
    elif name == "wilsonka": user = "Ken Wilson"
    elif name == "1119423": user = "Matthew Chu"
    return user

def create_header(doc, multiplier, pageType, multiImage):
    runlist=[] #list of things that should be in red
    #calculate reticle sizes
    oneWindow = (((5*float(stepRepeatX))+20)/10.0, ((5*float(stepRepeatY))+20)/10.0)
    reticleFive = (oneWindow[0]/100.0, oneWindow[1]/100.0)
    reticleOne = ((float(stepRepeatX)+4)/1000, (float(stepRepeatY)+4)/1000)
    windowSize = 101350
    productLine = "ASML"
    if pageType != "core" and pageType != "tert": productLine = pageType
    #get the structure
    structure = primary[:-1]
    if pageType == "tert": structure += "t"
    if pageType == "NUM": structure += "n"
    if pageType == "FS" or pageType == "FLIPPED": structure += "p_bcb"
    p = doc.add_paragraph(design+ "_0_revA.doc" + "\t\t\t\t" + date + "\t\t\t\t" + processType)
    p.add_run("\n\nRRFC")
    p.add_run("\n\nDevice: %s-%s" % (design[0:2].upper(), design[2:]))
    if pageType != "aeble": p.add_run("\nProduct: %dX %s" % (max(multiplier, 1), productLine))
    p.add_run("\nPO:")
    runlist.append(p.add_run("XXX"))
    p.add_run("\t\tRel: ")
    runlist.append(p.add_run("XXX"))
    user = get_user()
    if multiplier == 1 or multiplier == 5:
        p.add_run("\n\nCAD Layout Person: ")
        runlist.append(p.add_run(user))
    if pageType == "aeble": 
        for run in runlist:
            run.font.color.rgb = RGBColor(0xff, 0x0, 0x0)
        return
    p.add_run("\n\nSend File: ")
    if pageType == "FS": runlist.append(p.add_run("%sXXX.gds" % structure))
    else: runlist.append(p.add_run("%s.gds" % structure))
    if pageType == "core" and multiImage:
        p.add_run("\nStructure(s): ")
        runlist.append(p.add_run("P=%s, S=%s" % (primary, secondary)))
    elif pageType == "core":
        p.add_run("\nSTRUCTURE: ")
        runlist.append(p.add_run("(PRI) %s, (SEC) %s" % (primary, secondary)))
    elif pageType == "BS":
        p.add_run("\nSTRUCTURE: ")
        runlist.append(p.add_run("(PRI) %s, (SEC) %s" % (primary, secondary)))
        if multiImage:
            runlist.append(p.add_run(", (C) %s_center, (E) %s_edge" % (secondary, secondary)))
    elif pageType == "tert":
        p.add_run("\nSTRUCTURE: ")
        runlist.append(p.add_run("(TRI) %s" % structure))
    else:
        p.add_run("\nSTRUCTURE: ")
        runlist.append(p.add_run("%s" % structure))
    if multiplier == 1 or multiplier == 5:
        p.add_run("\n1X WINDOW: (-%6.2f, -%6.2f), (%6.2f, %6.2f)" % (oneWindow[0], oneWindow[1], oneWindow[0], oneWindow[1]))
    if multiplier == 5:
        p.add_run("\n5X RETICLE DATA WINDOW - (X) %5.3fMM, (Y) %5.3fMM" % (reticleFive[0], reticleFive[1]))
    elif multiplier == 1:
        p.add_run("\n1X RETICLE DATA WINDOW - (X) %5.3fMM, (Y) %5.3fMM" % (reticleOne[0], reticleOne[1]))
    else:
        runlist.append(p.add_run("\nDATA EXTENTS: NumBBllX,NumBBllY NumBBurX,NumBBurY XXX"))
        p.add_run("\nWINDOW SIZE (X) %d (Y) %d" % (windowSize, windowSize))
    if multiplier == 1 or multiplier == 5:
        p.add_run("\nS & R CENTERS - (X) %5.3fMM, (Y) %5.3fMM" % (float(stepRepeatX)/1000, float(stepRepeatY)/1000))
    for run in runlist: run.font.color.rgb = RGBColor(0xff, 0x0, 0x0)


def getLayers(mpssType, allLayers):
    result = []
    prevMpss = None
    for layer in allLayers:
        if layer.mpss == mpssType or (layer.mpss == "NA" and prevMpss == mpssType):
            result.append(layer)
        prevMpss = layer.mpss
    return result

def isTertiary(layer):
    return getAbbrev(layer.rTitle).split()[-1].strip() == "TRI"

def createFiveXTable(doc, layers, tableType, design):
    table = doc.add_table(rows = 1, cols = 7)
    hdr_cells = table.rows[0].cells
    headerList = ["ITEM NO.", "FILENAME", "GDS LAYER", "SUPPLIED DATA CD", "RETICLE CD LINE", "RETICLE TITLE", "BAR CODE"]
    runlist=[]
    for i in range(len(headerList)):
        run = hdr_cells[i].paragraphs[0].add_run(headerList[i])
        run.bold = True
        font = run.font
        font.size = Pt(8)
    prevLayer = None
    for layer in layers:
        runlist = []
        #filter tertiary layers out of the regular core page
        if layer.optional != "NA" and isTertiary(layer) and tableType != "tert":
            continue
        #ensure only tertiary layers go on the tertiary page, assumes all tertiary layers are also core layers
        elif layer.optional != "NA" and not isTertiary(layer) and tableType == "tert":
            continue
        #if its a layer to be merged
        if layer.mpss == "NA" and tableType != "tert":
            cells = table.add_row().cells
            runlist.append(cells[2].paragraphs[0].add_run(str(layer.gds)))
            for cell in cells: 
                cell.paragraphs[0].style.font.size = Pt(8)
        elif layer.mpss == "NA" and tableType == "tert": continue
        else:
            cells = table.add_row().cells
            runlist.append(cells[0].paragraphs[0].add_run(str(layer.itemNo)))
            #filename
            runlist.append(cells[1].paragraphs[0].add_run(build_filename(layer, design)))
            #GDS layer
            runlist.append(cells[2].paragraphs[0].add_run("%02d" % long(layer.gds)))
            #Supplied Data CD
            runlist.append(cells[3].paragraphs[0].add_run(layer.cad))
            #reticle CD line - grab only up to the +/- part
            runlist.append(cells[4].paragraphs[0].add_run(layer.rCD[:layer.rCD.index('m')+1]))
            #reticle title
            title = "%s-%s-00-%s %s" % (design[0:2].upper(), design[2:], layer.pLayer, getAbbrev(layer.rTitle))
            runlist.append(cells[5].paragraphs[0].add_run(title))
            #barcode
            pLayer = layer.pLayer.strip()
            if not pLayer.isnumeric(): pLayer = pLayer[:-1]
            while len(pLayer) < 3:
                pLayer = "0" + pLayer
            deviceType = design[:2].upper()
            barcode = "%s%sP00%s" % (get_device_type(deviceType).upper(), design[-3:], pLayer)
            abbrev = getAbbrev(layer.rTitle)
            if abbrev.split()[-1] == "SEC": barcode = barcode[:4] + "S" + barcode[5:]
            elif abbrev.split()[-1] == "TRI": barcode = barcode[:4] + "T" + barcode[5:]
            runlist.append(cells[6].paragraphs[0].add_run(barcode))
        #style formatting
        #make the font size 8 and turn all optional layers red
        for run in runlist:
            font = run.font
            font.size = Pt(8)
            if layer.optional == "O" or (layer.mpss == "NA" and prevLayer.optional == "O"):
                font.color.rgb = RGBColor(0xff, 0x0, 0x0)
        prevLayer = layer
    #manually sets column widths, needs to be done for both columns and cells to take effect
    table.autofit = False
    table.columns[0].width = Inches(0.7)
    for cell in table.columns[0].cells: cell.width = Inches(0.7)
    table.columns[1].width = Inches(0.88)
    for cell in table.columns[1].cells: cell.width = Inches(0.88)
    table.columns[2].width = Inches(0.63)
    for cell in table.columns[2].cells: cell.width = Inches(0.63)
    table.columns[3].width = Inches(0.75)
    for cell in table.columns[3].cells: cell.width = Inches(0.75)
    table.columns[4].width = Inches(0.94)
    for cell in table.columns[4].cells: cell.width = Inches(0.94)
    table.columns[5].width = Inches(1.43)
    for cell in table.columns[5].cells: cell.width = Inches(1.43)
    table.columns[6].width = Inches(0.76)
    for cell in table.columns[6].cells: cell.width = Inches(0.76)
    if tableType == "tert":
        table.columns[5].width = Inches(1.74)
        for cell in table.columns[5].cells: cell.width = Inches(1.74)
    #center align text in the table
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"

def get_device_type(deviceType) :
    if deviceType == "PD": leadingLetter = "D"
    elif deviceType == "LN": leadingLetter = "L"
    elif deviceType == "AM": leadingLetter = "M"
    elif deviceType == "PA": leadingLetter = "P"
    elif deviceType == "TR": leadingLetter = "R"
    elif deviceType == "PS": leadingLetter = "S"
    elif deviceType == "TS": leadingLetter = "T"
    elif deviceType == "SW": leadingLetter = "W"
    elif deviceType == "AT": leadingLetter = "A"
    else:
        print("Naming Convention Error. Unknown Device Type \"%s\"" % deviceType)
        leadingLetter = "X"
    return leadingLetter


def createOneXTable(doc, layers, design, mpss):
    table = doc.add_table(rows = len(layers)+1, cols = 6)
    headerList = ["ITEM NO.", "FILENAME", "GDS LAYER", "SUPPLIED DATA CD", "RETICLE CD LINE", "RETICLE TITLE"]
    hdr_cells = table.rows[0].cells
    for i in range(6):
        run = hdr_cells[i].paragraphs[0].add_run(headerList[i])
        run.bold = True
        font = run.font
        font.size = Pt(8)
    prevLayer = None
    for i in range(len(layers)):
        runlist = []
        layer = layers[i]
        cells = table.rows[i+1].cells
        #if its a layer that will be merged
        if layer.mpss == "NA":
            runlist.append(cells[2].paragraphs[0].add_run(str(layer.gds)))
            for cell in cells:
                cell.paragraphs[0].style.font.size = Pt(8)
        else:
            reticleLine = "See NOTES"
            title = None
            if layer.rCD == "NA" and mpss != "Backside": reticleLine = "NA"
            elif mpss != "Backside":
                reticleLine = layer.rCD[:layer.rCD.index('m')+1]
            if mpss == "Numbers": #the numbers reticle title is slightly different than the rest
                title = "%s-%s-00-%s" % (design[0:2].upper(), design[2:], getAbbrev(layer.rTitle))
            else:
                title = "%s-%s-00-%s %s" % (design[0:2].upper(), design[2:], layer.pLayer, getAbbrev(layer.rTitle))
            runlist.append(cells[0].paragraphs[0].add_run(str(layer.itemNo)))
            runlist.append(cells[1].paragraphs[0].add_run(build_filename(layer, design)))
            runlist.append(cells[2].paragraphs[0].add_run(str(layer.gds)))
            if mpss != "Backside":
                runlist.append(cells[3].paragraphs[0].add_run(layer.cad))
            else:
                runlist.append(cells[3].paragraphs[0].add_run("NA"))
            runlist.append(cells[4].paragraphs[0].add_run(reticleLine))
            runlist.append(cells[5].paragraphs[0].add_run(title))
        #style formatting
        #make the font size 8 and turn all optional layers red
        for run in runlist:
            font = run.font
            font.size = Pt(8)
            if layer.optional == "O" or (layer.mpss == "NA" and prevLayer.optional == "O"):
                font.color.rgb = RGBColor(0xff, 0x0, 0)
        prevLayer = layer
    #manually sets column widths, needs to be done for both columns and cells to take effect
    table.autofit = False
    table.columns[0].width = Inches(0.7)
    for cell in table.columns[0].cells: cell.width = Inches(0.7)
    table.columns[1].width = Inches(0.94)
    for cell in table.columns[1].cells: cell.width = Inches(0.94)
    table.columns[2].width = Inches(0.94)
    for cell in table.columns[2].cells: cell.width = Inches(0.94)
    table.columns[3].width = Inches(1.25)
    for cell in table.columns[3].cells: cell.width = Inches(1.25)
    table.columns[4].width = Inches(1.18)
    for cell in table.columns[4].cells: cell.width = Inches(1.18)
    table.columns[5].width = Inches(1.76)
    for cell in table.columns[5].cells: cell.width = Inches(1.76)
    #center align text in the table
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"

def digitized(layer, process):
    abbrev = getAbbrev(layer.rTitle)
    if abbrev == "NE" or abbrev == "GLASS": return "No"
    if process in GaN:
        if abbrev == "NE1" or abbrev == "MESA":
            return "No"
    return "Yes"

def deleteEmptyRows(table):
    for row in table.rows:
        empty = True
        for cell in row.cells:
            if cell.text != "":
                empty = False
                break
        if empty:
            table._tbl.remove(row._tr)



def createMultiImageTable(doc, layers, design, process):
    #table headers
    headerList = ["Item Number", "Field", "Use Data Structure", "GDS Number", "Supplied CD", "Reticle CD Line", "CD Digitized", "Reticle Title", "Barcode", "Center Data At"]
    table = doc.add_table(rows = (len(layers)*2)+1, cols = 10)
    hdr_cells = table.rows[0].cells
    for i in range(10):
        run = hdr_cells[i].paragraphs[0].add_run(headerList[i])
        run.bold = True
        font = run.font
        font.size = Pt(8)
    table.autofit = False
    #table content
    merging = False
    i = 0
    row = 1
    while i < len(layers):
        runlist = []
        if i == len(layers)-1: merging = False
        elif layers[i+1].mpss == "NA": merging = True
        layer = layers[i]
	#print layer.rTitle
        reticleLine = layer.rCD[:layer.rCD.index('m')+1]
        abbrev = getAbbrev(layer.rTitle)
        if abbrev == "SD PRI": abbrev = "SD"
        elif abbrev == "SD SEC":
            i += 2
            merging = False
            continue
        reticleTitle = "%s-%s-00-%s %s" % (design[0:2].upper(), design[2:], layer.pLayer, abbrev)
        pLayer = layer.pLayer.strip()
        if not pLayer.isnumeric(): pLayer = pLayer[:-1]
        while len(pLayer) < 3:
            pLayer = "0" + pLayer
	deviceType = design[:2].upper()
        barcode = "%s%sM00%s" % (get_device_type(deviceType).upper(), design[-3:], pLayer)
        #if abbrev.split()[-1] == "SEC": barcode = barcode[:4] + "S" + barcode[5:]
        #elif abbrev.split()[-1] == "TRI": barcode = barcode[:4] + "T" + barcode[5:]
        if not merging:
            runlist.append(table.cell(row, 0).paragraphs[0].add_run(str(layer.itemNo)))
            runlist.append(table.cell(row, 1).paragraphs[0].add_run("North"))
            runlist.append(table.cell(row+1, 1).paragraphs[0].add_run("South"))
            if isTertiary(layers[i]):
                runlist.append(table.cell(row, 2).paragraphs[0].add_run("T"))
                runlist.append(table.cell(row, 3).paragraphs[0].add_run(str(layer.gds)))
                runlist.append(table.cell(row, 4).paragraphs[0].add_run(layer.cad))
                runlist.append(table.cell(row, 5).paragraphs[0].add_run(reticleLine))
                runlist.append(table.cell(row, 6).paragraphs[0].add_run(digitized(layer, process)))
                runlist.append(table.cell(row, 7).paragraphs[0].add_run(reticleTitle))
                runlist.append(table.cell(row, 8).paragraphs[0].add_run(barcode))
                runlist.append(table.cell(row, 9).paragraphs[0].add_run("0, 29.75"))
                runlist.append(table.cell(row+1, 9).paragraphs[0].add_run("0, -29.75"))
                table.cell(row,0).merge(table.cell(row+1,0))
                table.cell(row,3).merge(table.cell(row+1,3))
                table.cell(row,4).merge(table.cell(row+1,4))
                table.cell(row,5).merge(table.cell(row+1,5))
                table.cell(row,6).merge(table.cell(row+1,6))
                table.cell(row,7).merge(table.cell(row+1,7))
                table.cell(row,8).merge(table.cell(row+1,8))
                row += 2
                i += 1
            else:
                runlist.append(table.cell(row, 2).paragraphs[0].add_run("P"))
                runlist.append(table.cell(row+1, 2).paragraphs[0].add_run("S"))
                runlist.append(table.cell(row, 3).paragraphs[0].add_run(str(layer.gds)))
                runlist.append(table.cell(row, 4).paragraphs[0].add_run(layer.cad))
                runlist.append(table.cell(row, 5).paragraphs[0].add_run(reticleLine))
                runlist.append(table.cell(row, 6).paragraphs[0].add_run(digitized(layer, process)))
                runlist.append(table.cell(row, 7).paragraphs[0].add_run(reticleTitle))
                runlist.append(table.cell(row, 8).paragraphs[0].add_run(barcode))
                runlist.append(table.cell(row, 9).paragraphs[0].add_run("0, 29.75"))
                runlist.append(table.cell(row+1, 9).paragraphs[0].add_run("0, -29.75"))
                table.cell(row,0).merge(table.cell(row+1,0))
                table.cell(row,3).merge(table.cell(row+1,3))
                table.cell(row,4).merge(table.cell(row+1,4))
                table.cell(row,5).merge(table.cell(row+1,5))
                table.cell(row,6).merge(table.cell(row+1,6))
                table.cell(row,7).merge(table.cell(row+1,7))
                table.cell(row,8).merge(table.cell(row+1,8))
                row += 2
                i += 1
        else:
            runlist.append(table.cell(row, 0).paragraphs[0].add_run("\n" + str(layer.itemNo)))
            runlist.append(table.cell(row, 1).paragraphs[0].add_run("North"))
            runlist.append(table.cell(row+2, 1).paragraphs[0].add_run("South"))
            runlist.append(table.cell(row, 2).paragraphs[0].add_run("P"))
            runlist.append(table.cell(row+2, 2).paragraphs[0].add_run("S"))
            runlist.append(table.cell(row, 3).paragraphs[0].add_run(str(layer.gds)))
            runlist.append(table.cell(row+1, 3).paragraphs[0].add_run(str(layers[i+1].gds)))
            runlist.append(table.cell(row+2, 3).paragraphs[0].add_run(str(layer.gds)))
            runlist.append(table.cell(row+3, 3).paragraphs[0].add_run(str(layers[i+1].gds)))
            runlist.append(table.cell(row, 4).paragraphs[0].add_run(layer.cad))
            runlist.append(table.cell(row+2, 4).paragraphs[0].add_run(layer.cad))
            runlist.append(table.cell(row, 5).paragraphs[0].add_run(reticleLine))
            runlist.append(table.cell(row+2, 5).paragraphs[0].add_run(reticleLine))
            runlist.append(table.cell(row, 6).paragraphs[0].add_run("\n" + digitized(layer, process)))
            runlist.append(table.cell(row, 7).paragraphs[0].add_run("\n" + reticleTitle))
            runlist.append(table.cell(row, 8).paragraphs[0].add_run("\n" + barcode))
            runlist.append(table.cell(row, 9).paragraphs[0].add_run("0, 29.75"))
            runlist.append(table.cell(row+2, 9).paragraphs[0].add_run("0, -29.75"))
            table.cell(row,0).merge(table.cell(row+3,0))
            table.cell(row,1).merge(table.cell(row+1,1))
            table.cell(row+2,1).merge(table.cell(row+3,1))
            table.cell(row,2).merge(table.cell(row+1,2))
            table.cell(row+2,2).merge(table.cell(row+3,2))
            table.cell(row,6).merge(table.cell(row+3,6))
            table.cell(row,7).merge(table.cell(row+3,7))
            table.cell(row,8).merge(table.cell(row+3,8))
            table.cell(row,9).merge(table.cell(row+1,9))
            table.cell(row+2,9).merge(table.cell(row+3,9))
            row += 4
            i += 2
        #style formatting
        #make the font size 8 and turn all optional layers red
        for run in runlist:
            font = run.font
            font.size = Pt(8)
            if layer.optional == "O":
                font.color.rgb = RGBColor(0xff, 0x0, 0)
        merging = False
    #manually sets column widths, needs to be done for both columns and cells to take effect
    table.columns[0].width = Inches(0.57)
    for cell in table.columns[0].cells: cell.width = Inches(0.57)
    table.columns[1].width = Inches(0.56)
    for cell in table.columns[1].cells: cell.width = Inches(0.56)
    table.columns[2].width = Inches(0.66)
    for cell in table.columns[2].cells: cell.width = Inches(0.66)
    table.columns[3].width = Inches(0.56)
    for cell in table.columns[3].cells: cell.width = Inches(0.56)
    table.columns[4].width = Inches(0.63)
    for cell in table.columns[4].cells: cell.width = Inches(0.63)
    table.columns[5].width = Inches(0.58)
    for cell in table.columns[5].cells: cell.width = Inches(0.58)
    table.columns[6].width = Inches(0.58)
    for cell in table.columns[6].cells: cell.width = Inches(0.58)
    table.columns[7].width = Inches(1.41)
    for cell in table.columns[7].cells: cell.width = Inches(1.41)
    table.columns[8].width = Inches(0.81)
    for cell in table.columns[8].cells: cell.width = Inches(0.81)
    table.columns[9].width = Inches(0.61)
    for cell in table.columns[9].cells: cell.width = Inches(0.61)
    #center align text in the table
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(8)
    deleteEmptyRows(table)


def getAebleMap(doc, design, multi, pathdir):
#    user = subprocess.Popen(["whoami"], stdout = subprocess.PIPE).communicate()[0][:-1]
#    path = "/users/%s/cad_layout_61/%s/release/png/" % (user, design)
    path = pathdir + "/png/" 
    for fname in os.listdir(path):
        if "aeblemap" in fname:
            if multi:
                #add in notes
                letters = design[:2]
                numbers = design[2:]
                while numbers[0] == "0": numbers = numbers[1:]
                prime = letters + numbers + "_0p"
                second = letters + numbers + "_0s"
                edge = second + "_edge"
                center = second + "_center"
                runlist = []
                p = doc.add_paragraph()
                runlist.append(p.add_run ("Special Notes-"))
                p.add_run("\n-  On Aeblemap below please place the %s structure in the two locations designated with the" % edge)
                runlist.append(p.add_run(" E"))
                p.add_run("\n-  On Aeblemap below please place the %s structure in the one location designated with the" % center)
                runlist.append(p.add_run(" C"))
                p.add_run(" (center of wafer)")
                p.add_run("\n-  On Aeblemap below please place the %s structure in the" % second)
                runlist.append(p.add_run(" 6 "))
                p.add_run("remaining Secondary locations")
                p.add_run("\n-  On Aeblemap below place the %s structure in all of the Primary locations" % prime)
                for run in runlist:
                    run.bold = True
            #add aeblemap under it
            os.system("cp %s%s tempAeble.png" % (path, fname))
            doc.add_picture("tempAeble.png", width = Inches(6))
            os.system("rm tempAeble.png")
            return True
        else:
            print("ERROR: Aeble map image not found in directory: " + path + "\n")
    return False


########################################################################################################
################################START OF SCRIPT#########################################################
########################################################################################################

#takes a commandline argument, the path to the photo file
photofile = sys.argv[1]

#parse the photofile
photofile = open(photofile, "r")
lines = photofile.readlines()
photofile.close()
design = lines[0].split()[-1]
primary = lines[1].split()[-1]
secondary = lines[2].split()[-1]
processType = lines[3].split()[-1]
stepRepeatX = lines[4].split()[-2]
stepRepeatY = lines[4].split()[-1]

#strip out the letter for all except P80, ie P46H becomes P46
if not processType[-1].isdigit() and not ("P80" in processType):
    strippedProcess = processType[:-1]
else: strippedProcess = processType

#parse the RSS
layerList = process_worksheet(strippedProcess)

#determine if its to be a multi image mask or not
multi = False
if sys.argv[2] == "multi":
    multi = True

#create the MPSS document
doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
#get the date(for header)
today = datetime.date.today()
date = today.strftime("%m/%d/%Y")


##############################################cores##########################################################
create_header(doc, 5, "core", multi)

coreLayers = getLayers("Core", layerList)
#for lay in coreLayers:
#    print lay.rTitle

if multi:
    createMultiImageTable(doc, coreLayers, design, strippedProcess)
    #add the north/south image and center align it
    doc.add_picture(northSouthDiagram)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else: 
    createFiveXTable(doc, coreLayers, "core", design)

doc.add_paragraph("NOTES:")
doc.add_paragraph("END NOTES")
doc.add_page_break()

######################################################tertiary###############################################
if strippedProcess in GaAs and not multi:
    create_header(doc, 5, "tert", multi)

    createFiveXTable(doc, coreLayers, "tert", design)

    doc.add_paragraph("NOTES:")
    doc.add_paragraph("END NOTES:")
    doc.add_page_break()

#####################################################backside#############################################
bsLayers = getLayers("Backside", layerList)

if len(bsLayers) > 0:
    create_header(doc, 1, "BS", multi)

    createOneXTable(doc, bsLayers, design, "Backside")

    #grab the grid and via layers for use in the notes section
    via = None
    grid = None
    sstop = None
    for layer in bsLayers:
        if layer.rTitle == "VIA": via = layer
        if layer.rTitle == "GRID": grid = layer
        if layer.rTitle == "SOLDERSTOP": sstop = layer

    doc.add_paragraph("NOTES:")
    p = doc.add_paragraph("""DATA: Right reading chrome side up

TITLES: Right reading chrome side up

"CD" Measurements performed on Photronics closure patters are acceptable for masks listed on this page

Shift around the array generation by waferShiftX and waferShiftY found on the Aeblemap""")

    if via != None and via.pLayer != "NA" and grid != None and grid.pLayer != "NA":
        if sstop != None and sstop.pLayer != "NA": 
            p.add_run("\n\nPlease shift the array of %s and %s so that they overlay with the %s" % (sstop.pLayer, grid.pLayer, via.pLayer))
        else: 
            p.add_run("\n\nPlease shift the array of %s so that it overlays with the %s" % (grid.pLayer, via.pLayer))
    doc.add_paragraph("END NOTES")
    doc.add_page_break()

################################################aeble map########################################################
create_header(doc, 5, "aeble", multi)

run = doc.add_paragraph().add_run("%s-%s Aeblemap" % (design[0:2].upper(), design[2:]))
run.font.size = Pt(20)
run.bold = True
pathdir = os.path.dirname(sys.argv[1])
if not getAebleMap(doc, design, multi, pathdir) :
    run = doc.add_paragraph().add_run("PUT AEBLE MAP HERE XXX")
    run.font.color.rgb = RGBColor(0xff, 0x0, 0x0)

doc.add_page_break()

##################################################numbers page####################################################
numLayers  = getLayers("Numbers", layerList)
removed = None
for layer in numLayers:
    if str(layer.gds) == "37" and multi:
        removed = layer
    elif str(layer.gds) == "55" and not multi: 
        removed = layer
if removed != None: numLayers.remove(removed)
if len(numLayers) > 0:
    create_header(doc, 0, "NUM", multi)

    createOneXTable(doc, numLayers, design, "Numbers")

    doc.add_paragraph("NOTES:")
    doc.add_paragraph(
    """DATA: Wrong reading chrome side up

TITLES: Wrong reading chrome side up

Please center data on the mask

All 1X CD crosses are located in the four corners of the data extents""")
    doc.add_paragraph("END NOTES")
    doc.add_page_break()

###############################################frontside################################################################
fsLayers = getLayers("Frontside", layerList)
if len(fsLayers) > 0:
    create_header(doc, 0, "FS", multi)

    createOneXTable(doc, fsLayers, design, "Frontside")

    doc.add_paragraph("NOTES:")
    doc.add_paragraph(
    """DATA: Wrong reading chrome side up

TITLES: Wrong reading chrome side up

Please center data on the mask

All 1X CD crosses are located in the four corners of the data extents""")
    doc.add_paragraph("END NOTES")

    doc.add_page_break()

#############################################flipped######################################################################
flippedLayers = getLayers("Flipped", layerList)

if len(flippedLayers) > 0:
    create_header(doc, 0, "FLIPPED", multi)

    createOneXTable(doc, flippedLayers, design, "Flipped")

    doc.add_paragraph("NOTES:")
    doc.add_paragraph(
    """DATA: Wrong reading chrome side up

TITLES: Wrong reading chrome side up

Please center data on the mask

All 1X CD crosses are located in the four corners of the data extents""")
    doc.add_paragraph("END NOTES")

#save and close
docName = "%s_0_revA.docx" % design
#docxPath = "~/cad_layout_61/%s/release/" % design
doc.save(docName)
os.system("mv %s %s" % (docName, pathdir))
sys.stdout.write("MPSS complete, check release folder\n")
