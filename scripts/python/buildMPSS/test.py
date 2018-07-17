# Command Line call
# Design Name Error
/net/nfs-n501/ANFS01/apps/python/python27/bin/python /users/1127110/project/python/buildMPSS/trunk/buildMPSS.py "/users/1127110/cad_layout_61/p80b_wafer_map/release/p80b_wafer_map.photo" "single"
# Fixed design name
/net/nfs-n501/ANFS01/apps/python/python27/bin/python /users/1127110/project/python/buildMPSS/trunk/buildMPSS.py "/users/1127110/cad_layout_61/am105_0p/release/am105_0p.photo" "single"
# Best design name 
/net/nfs-n501/ANFS01/apps/python/python27/bin/python /users/1127110/project/python/buildMPSS/trunk/buildMPSS.py "/users/1127110/cad_layout_61/am0105/release/am0105.photo" "single"
# Best design name 
/net/nfs-n501/ANFS01/apps/python/python27/bin/python /users/1127110/project/python/buildMPSS/trunk/buildMPSS.py "/scratch/oaConversion/oaConversion/designs/oaLibs/pa0715/release/pa0715.photo" "single"

# interactive
# test .photo file

import sys
import os
import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import subprocess

warnings.filterwarnings("ignore")
#redirect error messages to stdout so cadence can read them
sys.stderr = sys.stdout

# test .photo file
photofile = "/users/1127110/cad_layout_61/p80b_wafer_map/release/p80b_wafer_map.photo"

# test multi-image mask
multi = False
tableType = "core"
layers = coreLayers

for layer in layers:
        runlist = []
        #filter tertiary layers out of the regular core page
        if layer.optional != "NA" and isTertiary(layer) and tableType != "tert":
            continue
        #ensure only tertiary layers go on the tertiary page, assumes all tertiary layers are also core layers
        elif layer.optional != "NA" and not isTertiary(layer) and tableType == "tert":
            continue
        #if its a layer to be merged
        if layer.mpss == "NA" and tableType != "tert":
            cells = table.add_row().cells
            runlist.append(cells[2].paragraphs[0].add_run(str(layer.gds)))
            for cell in cells: 
                cell.paragraphs[0].style.font.size = Pt(8)
        elif layer.mpss == "NA" and tableType == "tert": continue
        else:
            cells = table.add_row().cells
            runlist.append(cells[0].paragraphs[0].add_run(str(layer.itemNo)))
            #filename
            runlist.append(cells[1].paragraphs[0].add_run(build_filename(layer, design)))
            #GDS layer
            runlist.append(cells[2].paragraphs[0].add_run("%02d" % long(layer.gds)))
            #Supplied Data CD
            runlist.append(cells[3].paragraphs[0].add_run(layer.cad))
            #reticle CD line - grab only up to the +/- part
            runlist.append(cells[4].paragraphs[0].add_run(layer.rCD[:layer.rCD.index('m')+1]))
            #reticle title
            title = "%s-%s-00-%s %s" % (design[0:2].upper(), design[2:], layer.pLayer, getAbbrev(layer.rTitle))
            runlist.append(cells[5].paragraphs[0].add_run(title))
            #barcode
            pLayer = layer.pLayer.strip()
            if not pLayer.isnumeric(): pLayer = pLayer[:-1]
            while len(pLayer) < 3:
                pLayer = "0" + pLayer
            deviceType = design[:2].upper()
            barcode = "%s%sP00%s" % (get_device_type(deviceType).upper(), design[-3:], pLayer)
            abbrev = getAbbrev(layer.rTitle)
            if abbrev.split()[-1] == "SEC": barcode = barcode[:4] + "S" + barcode[5:]
            elif abbrev.split()[-1] == "TRI": barcode = barcode[:4] + "T" + barcode[5:]
            runlist.append(cells[6].paragraphs[0].add_run(barcode))

